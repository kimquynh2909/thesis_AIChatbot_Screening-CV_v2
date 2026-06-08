from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from src.preprocessing.text_cleaner import clean_for_display
from src.utils.file_utils import read_text_file


@dataclass
class ParsedDocument:
    filename: str
    extension: str
    raw_text: str
    cleaned_text: str


class DocumentParser:
    """Extract text from PDF, DOCX, and TXT files used in resumes and JDs."""

    def parse(self, source: str | Path | bytes | BinaryIO, filename: str | None = None) -> ParsedDocument:
        if isinstance(source, (str, Path)):
            path = Path(source)
            file_bytes = path.read_bytes()
            resolved_filename = filename or path.name
        elif isinstance(source, bytes):
            file_bytes = source
            resolved_filename = filename or "uploaded_document.txt"
        else:
            if hasattr(source, "getvalue"):
                file_bytes = source.getvalue()
            else:
                file_bytes = source.read()
            resolved_filename = filename or getattr(source, "name", "uploaded_document.txt")

        extension = Path(resolved_filename).suffix.lower()
        if extension == ".pdf":
            raw_text = self._parse_pdf(file_bytes)
        elif extension == ".docx":
            raw_text = self._parse_docx(file_bytes)
        elif extension == ".txt" or not extension:
            raw_text = self._parse_txt(file_bytes)
            extension = ".txt"
        else:
            raise ValueError(f"Unsupported file type '{extension}'. Supported types are PDF, DOCX, and TXT.")

        cleaned = clean_for_display(raw_text)
        return ParsedDocument(
            filename=resolved_filename,
            extension=extension,
            raw_text=raw_text,
            cleaned_text=cleaned,
        )

    def parse_path(self, path: Path) -> ParsedDocument:
        if path.suffix.lower() == ".txt":
            raw = read_text_file(path)
            return ParsedDocument(path.name, ".txt", raw, clean_for_display(raw))
        return self.parse(path)

    @staticmethod
    def _parse_txt(file_bytes: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode(errors="ignore")

    @staticmethod
    def _parse_docx(file_bytes: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("python-docx is required to parse DOCX files.") from exc

        document = Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        table_text: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))
        return "\n".join(paragraphs + table_text)

    @staticmethod
    def _parse_pdf(file_bytes: bytes) -> str:
        errors: list[str] = []

        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(BytesIO(file_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                return text
        except Exception as exc:  # PDF parsers often fail on malformed PDFs.
            errors.append(f"PyPDF2: {exc}")

        try:
            import pdfplumber

            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                if text.strip():
                    return text
        except Exception as exc:
            errors.append(f"pdfplumber: {exc}")

        try:
            import fitz

            document = fitz.open(stream=file_bytes, filetype="pdf")
            text = "\n".join(page.get_text("text") for page in document)
            if text.strip():
                return text
        except Exception as exc:
            errors.append(f"PyMuPDF: {exc}")

        joined = "; ".join(errors) if errors else "no parser produced text"
        raise ValueError(f"Could not extract text from PDF: {joined}")
